import streamlit as st
import pdfplumber
import pandas as pd
from docx import Document
import io

st.set_page_config(page_title="PDF Converter", page_icon="📄", layout="centered")

st.title("📄 PDF to Excel / Word Converter")

uploaded_file = st.file_uploader("Upload your PDF file", type=["pdf"])
output_format = st.selectbox("Choose output format", ["Excel", "Word"])

if uploaded_file and st.button("Convert"):
    with pdfplumber.open(uploaded_file) as pdf:
        if output_format == "Excel":
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
                for page_num, page in enumerate(pdf.pages, start=1):
                    tables = page.extract_tables()
                    for idx, table in enumerate(tables, start=1):
                        if table:
                            df = pd.DataFrame(table)
                            sheet_name = f"Page{page_num}_Tbl{idx}"[:31]
                            df.to_excel(writer, index=False, sheet_name=sheet_name)
            buffer.seek(0)
            st.download_button(
                label="⬇️ Download Excel File",
                data=buffer,
                file_name="converted_output.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

        elif output_format == "Word":
            doc = Document()
            for page_num, page in enumerate(pdf.pages, start=1):
                doc.add_heading(f"Page {page_num}", level=1)
                tables = page.find_tables()
                words = page.extract_words(use_text_flow=True)
                filtered_words = []
                for w in words:
                    wx0, wy0, wx1, wy1 = w["x0"], w["top"], w["x1"], w["bottom"]
                    inside_table = False
                    for t in tables:
                        x0, y0, x1, y1 = t.bbox
                        if (wx0 >= x0 and wx1 <= x1 and wy0 >= y0 and wy1 <= y1):
                            inside_table = True
                            break
                    if not inside_table:
                        filtered_words.append(w)

                filtered_words.sort(key=lambda x: (x['top'], x['x0']))
                lines, current_line, current_y = [], [], None
                for w in filtered_words:
                    if current_y is None or abs(w['top'] - current_y) < 5:
                        current_line.append(w["text"])
                        current_y = w["top"]
                    else:
                        lines.append(" ".join(current_line))
                        current_line = [w["text"]]
                        current_y = w["top"]
                if current_line:
                    lines.append(" ".join(current_line))

                for line in lines:
                    doc.add_paragraph(line)

                for table in tables:
                    extracted = table.extract()
                    if extracted:
                        doc.add_paragraph()
                        t = doc.add_table(rows=len(extracted), cols=len(extracted[0]))
                        t.style = "Table Grid"
                        for r in range(len(extracted)):
                            for c in range(len(extracted[0])):
                                t.cell(r, c).text = extracted[r][c] or ""
                        doc.add_paragraph()
                doc.add_page_break()

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="⬇️ Download Word File",
                data=buffer,
                file_name="converted_output.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
